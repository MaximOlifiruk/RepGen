from django.shortcuts import render
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image
import json
import base64
from io import BytesIO
import pytesseract
import os


def index(request):
    return render(request, "main/start.html")


def main_page(request):
    return render(request, "main/index.html")
    

def post_data_title(request):
    temp = {}

    temp['student_surname'] = request.POST.get('studentSurname')
    temp['student_name'] = request.POST.get('studentName')
    temp['student_patronymic'] = request.POST.get('studentPatronymic')
    temp['teacher_surname'] = request.POST.get('teacherSurname')
    temp['teacher_name'] = request.POST.get('teacherName')
    temp['teacher_patronymic'] = request.POST.get('teacherPatronymic')
    temp['number_lab'] = request.POST.get('numberLab')
    temp['name_lab'] = request.POST.get('nameLab')
    temp['year_lab'] = request.POST.get('yearLab')
    temp['paragraphs_data'] = json.loads(request.POST.get('paragraphsData'))

    create_document(temp)
    return render(request, "main/end.html")


def create_paragraph(document, paragraph_name):
    heading_paragraph = document.add_paragraph()

    name_run = heading_paragraph.add_run(paragraph_name)
    name_run.bold = True
    name_run.font.size = Pt(16)

    heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def create_paragraph_text(document, paragraph_text):
    text_paragraph = document.add_paragraph()

    text_run = text_paragraph.add_run(paragraph_text)
    text_run.font.size = Pt(14)

    text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def create_paragraph_image(document, paragraph_image):
    image_data = paragraph_image
    header, encoded_data = image_data.split(',', 1)

    decoded_data = base64.b64decode(encoded_data)
    image_stream = BytesIO(decoded_data)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Cm(6.8), height=Cm(6.8))

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def create_paragraph_text_from_image(document, paragraph_image_text):
    header, encoded_data = paragraph_image_text.split(',', 1)

    decoded_data = base64.b64decode(encoded_data)
    image_stream = BytesIO(decoded_data)
    image = Image.open(image_stream)

    image_path = "temp_image.png"
    image.save(image_path)
    string = pytesseract.image_to_string(image_path)

    os.remove(image_path)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(string)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def create_document(temp):
    document = DocxTemplate('base_report/base_report.docx')
    document.render(temp)

    normal_style = document.styles['Normal']
    font = normal_style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    for paragraph_data in temp['paragraphs_data']:
        paragraph_name = paragraph_data['paragraphName']
        paragraph_content = paragraph_data['paragraphContent']

        create_paragraph(document, paragraph_name)

        for content in paragraph_content:
            if content['number'] == 1:
                create_paragraph_text(document, content['value'])
            elif content['number'] == 2:
                create_paragraph_image(document, content['value'])
            elif content['number'] == 3:
                create_paragraph_text_from_image(document, content['value'])

    document.save("result_report/result_document.docx")


